from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "DEPLOYMENT_GUIDE.md"
OUTPUT = ROOT / "ASR_Complete_Deployment_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
CODE_FILL = "F3F5F7"
MUTED = RGBColor(90, 99, 110)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths_inches)
    table.width = Inches(total)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(total * 1440)))
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("ASR Global Solutions  |  Deployment Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Platform-neutral deployment reference  |  August 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(115)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("DEPLOYMENT MANUAL")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("ASR AI Commerce Engagement")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(29)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(46)
    run = subtitle.add_run("Complete step-by-step deployment guide for static hosting, PaaS, VPS, and containers")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=3, cols=2)
    values = (("Audience", "Developers and hosting administrators"), ("Applies to", "Any provider supporting static files, Node.js, Python, or containers"), ("Version", "1.0  |  August 2026"))
    for row, values_row in zip(meta.rows, values):
        for cell, value in zip(row.cells, values_row):
            cell.text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row.cells[0], LIGHT_BLUE)
    set_table_geometry(meta, [1.4, 5.1])

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(34)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Important: static hosting cannot run the quotation, AI, tracking, email, or storage APIs.")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(156, 87, 0)

    doc.add_page_break()


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shading)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)


def add_markdown_table(doc, rows):
    parsed = [[item.strip() for item in row.strip().strip("|").split("|")] for row in rows]
    headers = parsed[0]
    body = parsed[2:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for values in body:
        cells = table.add_row().cells
        for index, text in enumerate(values):
            cells[index].text = text
            for run in cells[index].paragraphs[0].runs:
                run.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    column_count = len(headers)
    widths = [1.45, 0.7, 0.7, 0.7, 0.7, 2.25] if column_count == 6 else [6.5 / column_count] * column_count
    set_table_geometry(table, widths)


def add_inline_markup(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def add_body_from_markdown(doc):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 1  # The DOCX cover replaces the Markdown title.
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|---"):
            table_rows = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                table_rows.append(lines[index])
                index += 1
            add_markdown_table(doc, table_rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            doc.add_heading(heading.group(2), level=level)
        elif re.match(r"^\d+\.\s+", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_markup(paragraph, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markup(paragraph, line[2:])
        elif line.strip():
            paragraph = doc.add_paragraph()
            add_inline_markup(paragraph, line)
        index += 1


def build():
    doc = Document()
    style_document(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc)
    add_body_from_markdown(doc)
    doc.core_properties.title = "ASR AI Commerce Engagement - Complete Deployment Guide"
    doc.core_properties.subject = "Platform-neutral deployment instructions"
    doc.core_properties.author = "ASR Global Solutions"
    doc.core_properties.keywords = "ASR, deployment, Node.js, Python, VPS, PaaS, Docker"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
